import PyPDF2
import re
from docx import Document


def extract_text_from_pdf(file_path):
    """
    Extracts text from a PDF file.
    """
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
    return text


def extract_text_from_word(file_path):
    """
    Extracts text from a Word (.docx) document.
    """
    try:
        doc = Document(file_path)
        return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    except Exception as e:
        print(f"An error occurred while extracting text from the Word document: {e}")
        return ""


def extract_hyperlinks_from_pdf(file_path):
    """
    Extracts hyperlinks from a PDF file.
    """
    hyperlinks = []
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            if '/Annots' in page:
                annotations = page['/Annots']
                for annotation in annotations:
                    annotation_obj = annotation.get_object()
                    if '/A' in annotation_obj and '/URI' in annotation_obj['/A']:
                        uri = annotation_obj['/A']['/URI']
                        hyperlinks.append(uri)
    return hyperlinks


def extract_hyperlinks_from_word(file_path):
    """
    Extracts hyperlinks from a Word (.docx) document.
    """
    hyperlinks = []
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            for run in paragraph._element.xpath('.//w:hyperlink//w:r'):
                for parent in run.xpath('ancestor::w:hyperlink'):
                    if parent.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id'):
                        rel_id = parent.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        target = doc.part.rels[rel_id].target_ref
                        hyperlinks.append(target)
        return hyperlinks
    except Exception as e:
        print(f"An error occurred while extracting hyperlinks from the Word document: {e}")
        return []


def remove_large_digit_strings_and_special_chars(text):
    """
    Removes lines containing 5 or more digits or the '@' character,
    but preserves lines with exactly 4 digits (like years).
    """
    def has_five_or_more_digits(s):
        digit_sequences = re.findall(r'\d+', s)
        return any(len(seq) >= 5 for seq in digit_sequences)

    def extract_links(s):
        url_pattern = r'(https?://\S+|www\.\S+)'
        return re.findall(url_pattern, s)

    lines = text.split('\n')
    deleted_lines = [line for line in lines if has_five_or_more_digits(line) or '@' in line]
    filtered_lines = [line for line in lines if not has_five_or_more_digits(line) and '@' not in line]

    links = []
    for line in lines:
        links.extend(extract_links(line))

    print("Extracted Links:", links)

    return '\n'.join(filtered_lines), deleted_lines  # Return both filtered text and deleted lines