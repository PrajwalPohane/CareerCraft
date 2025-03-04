import os
import re
from pathlib import Path
import PyPDF2
from docx import Document
from fastapi import FastAPI, HTTPException, UploadFile, File, APIRouter
from pydantic import BaseModel
from typing import List, Dict
from groq import Groq
import tempfile
import json
import io

app = FastAPI()

# Initialize Groq API Client
api_key = "YOUR_GROQ_API_KEY"  # Replace with your actual API key
client = Groq(api_key=api_key)

def extract_text_from_pdf(file_content):
    """Extracts text from a PDF file content."""
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error extracting PDF text: {str(e)}")
        raise Exception(f"PDF extraction failed: {str(e)}")

def extract_text_from_docx(file_content):
    """Extracts text from a Word document content."""
    try:
        doc = Document(io.BytesIO(file_content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error extracting DOCX text: {str(e)}")
        raise Exception(f"DOCX extraction failed: {str(e)}")

def extract_hyperlinks_from_pdf(file_path):
    """Extracts hyperlinks from a PDF file."""
    hyperlinks = []
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            if '/Annots' in page:
                for annotation in page['/Annots']:
                    annotation_obj = annotation.get_object()
                    if '/A' in annotation_obj and '/URI' in annotation_obj['/A']:
                        hyperlinks.append(annotation_obj['/A']['/URI'])
    return hyperlinks

def extract_hyperlinks_from_word(file_path):
    """Extracts hyperlinks from a Word (.docx) document."""
    hyperlinks = []
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            for run in paragraph._element.xpath('.//w:hyperlink//w:r'):
                for parent in run.xpath('ancestor::w:hyperlink'):
                    rel_id = parent.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    if rel_id:
                        hyperlinks.append(doc.part.rels[rel_id].target_ref)
    except Exception as e:
        print(f"Error extracting hyperlinks from Word document: {e}")
    return hyperlinks

def remove_large_digit_strings_and_special_chars(text):
    """Removes lines with 5+ digits or '@', preserving 4-digit (year) lines."""
    def has_five_or_more_digits(s):
        return any(len(seq) >= 5 for seq in re.findall(r'\d+', s))

    def extract_links(s):
        return re.findall(r'(https?://\S+|www\.\S+)', s)

    lines = text.split('\n')
    deleted_lines = [line for line in lines if has_five_or_more_digits(line) or '@' in line]
    filtered_lines = [line for line in lines if line not in deleted_lines]

    links = [link for line in lines for link in extract_links(line)]
    print("Extracted Links:", links)
    
    return '\n'.join(filtered_lines), deleted_lines, links

router = APIRouter()

class TextAnalyzer:
    def __init__(self):
        self.client = client

    def analyze_text(self, text):
        """Analyzes text and determines if it's a resume or job description."""
        try:
            # First, determine if it's a resume or job description
            detection_prompt = f"""
            Analyze this text and determine if it's from a RESUME or a JOB DESCRIPTION.
            Respond with only one word: "RESUME" or "JD".
            
            Text to analyze:
            {text[:1000]}
            """
            
            detection_response = self.client.chat.completions.create(
                messages=[
                    {"role": "system", "content": "You are a resume analyzer. Classify the text as RESUME or JD."},
                    {"role": "user", "content": detection_prompt}
                ],
                model="llama3-70b-8192",
                temperature=0.1
            )
            
            text_type = detection_response.choices[0].message.content.strip()
            
            # Now parse the content based on the type
            if text_type == "RESUME":
                parse_prompt = f"""
                You are a resume parser. Extract information from the resume and return it in the exact JSON format shown below.
                Do not include any other text or explanation, only return the JSON object.

                Required JSON format:
                {{
                    "Name": "Candidate Name",
                    "Experience": [
                        {{
                            "title": "Job Title",
                            "company": "Company Name",
                            "duration": "Duration",
                            "responsibilities": ["responsibility 1", "responsibility 2"]
                        }}
                    ],
                    "Education": [
                        {{
                            "degree": "Degree Name",
                            "institution": "Institution Name",
                            "year": "Year"
                        }}
                    ],
                    "Skills": ["skill1", "skill2", "skill3"],
                    "Projects": [
                        {{
                            "name": "Project Name",
                            "description": "Project Description",
                            "technologies": ["tech1", "tech2"]
                        }}
                    ],
                    "Certifications": ["certification1", "certification2"],
                    "Contact": {{
                        "email": "email address",
                        "phone": "phone number",
                        "location": "location"
                    }}
                }}

                Resume text:
                {text}
                """
            else:
                parse_prompt = f"""
                You are a job description parser. Extract information from the job description and return it in the exact JSON format shown below.
                Do not include any other text or explanation, only return the JSON object.

                Required JSON format:
                {{
                    "JobTitle": "Title of the position",
                    "Company": "Company name",
                    "Requirements": ["requirement1", "requirement2"],
                    "Responsibilities": ["responsibility1", "responsibility2"],
                    "QualificationsRequired": ["qualification1", "qualification2"],
                    "PreferredSkills": ["skill1", "skill2"]
                }}

                Job Description text:
                {text}
                """

            parse_response = self.client.chat.completions.create(
                messages=[
                    {
                        "role": "system",
                        "content": "You are a document parser. Return only valid JSON without any additional text or markdown formatting."
                    },
                    {
                        "role": "user",
                        "content": parse_prompt
                    }
                ],
                model="llama3-70b-8192",
                temperature=0.3  # Lower temperature for more consistent output
            )

            # Clean and parse the response
            parsed_text = parse_response.choices[0].message.content.strip()
            # Remove any potential markdown or explanation
            if '```json' in parsed_text:
                parsed_text = re.search(r'```json\s*(.*?)\s*```', parsed_text, re.DOTALL)
                if parsed_text:
                    parsed_text = parsed_text.group(1)
            elif '```' in parsed_text:
                parsed_text = re.search(r'```\s*(.*?)\s*```', parsed_text, re.DOTALL)
                if parsed_text:
                    parsed_text = parsed_text.group(1)
            
            try:
                parsed_info = json.loads(parsed_text)
                # Validate required fields are present
                if text_type == "RESUME":
                    required_fields = ["Name", "Experience", "Education", "Skills"]
                else:
                    required_fields = ["JobTitle", "Requirements", "Responsibilities"]
                
                missing_fields = [field for field in required_fields if field not in parsed_info]
                if missing_fields:
                    raise Exception(f"Missing required fields: {', '.join(missing_fields)}")
                
                return parsed_info, text_type
            except json.JSONDecodeError as e:
                print(f"JSON parsing error: {str(e)}")
                print(f"Raw response: {parsed_text}")
                raise Exception("Failed to parse LLM response as JSON. Please try again.")

        except Exception as e:
            print(f"Error in analyze_text: {str(e)}")
            raise Exception(f"Text analysis failed: {str(e)}")

def process_document(file_path):
    """Processes a document (PDF or Word) and extracts relevant information."""
    file_extension = Path(file_path).suffix.lower()
    
    if file_extension == '.pdf':
        raw_text = extract_text_from_pdf(file_path)
        hyperlinks = extract_hyperlinks_from_pdf(file_path)
    elif file_extension == '.docx':
        raw_text = extract_text_from_docx(file_path)
        hyperlinks = extract_hyperlinks_from_word(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")
    
    filtered_text, deleted_lines, extracted_links = remove_large_digit_strings_and_special_chars(raw_text)
    return filtered_text, hyperlinks + extracted_links, deleted_lines, raw_text

@router.post("/upload_resume")
async def upload_resume(file: UploadFile = File(...)):
    try:
        # Read file content
        file_content = await file.read()
        
        # Extract text based on file type
        file_extension = Path(file.filename).suffix.lower()
        
        if file_extension == '.pdf':
            text = extract_text_from_pdf(file_content)
        elif file_extension in ['.docx', '.doc']:
            text = extract_text_from_docx(file_content)
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file format. Please upload PDF or DOCX files only."
            )

        # Analyze the text
        analyzer = TextAnalyzer()
        parsed_info, text_type = analyzer.analyze_text(text)
        
        return {
            "document_type": text_type,
            "parsed_information": parsed_info,
            "raw_text": text
        }
            
    except Exception as e:
        print(f"Upload resume error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=str(e)
        )

# Run the FastAPI application
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
